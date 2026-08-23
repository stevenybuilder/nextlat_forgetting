#!/usr/bin/env python
"""Render report/blog.md -> report/NextLat_Predictive_Geometry.docx.

The markdown file is the single source of truth for the writeup. Live numbers are
injected from results/ before rendering, so the Word doc never drifts from the
artifacts: a `{{live:KEY}}` token in the markdown is replaced with the value of KEY
in results/live_numbers.json, or with `[pending]` if that key has not been measured
yet. A rendered `[pending]` is a feature -- it makes an unfinished claim visible in
the document instead of letting prose quietly outrun the evidence.
"""
import json
import pathlib
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

ROOT = pathlib.Path(__file__).resolve().parent.parent
SRC = ROOT / "report" / "blog.md"
OUT = ROOT / "report" / "NextLat_Predictive_Geometry.docx"
LIVE = ROOT / "results" / "live_numbers.json"

MONO = "Menlo"
BODY = "Charter"
PENDING = "[pending]"


def load_live():
    if LIVE.exists():
        return json.loads(LIVE.read_text())
    return {}


def substitute(text, live):
    missing = set()

    def repl(m):
        key = m.group(1).strip()
        if key in live:
            return str(live[key])
        missing.add(key)
        return PENDING

    return re.sub(r"\{\{live:([^}]+)\}\}", repl, text), missing


def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.35
    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        st = doc.styles[name]
        st.font.name = BODY
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        st.font.bold = True


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = MONO
    run.font.size = Pt(9)


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [r for r in cells if not all(set(c) <= set("-: ") for c in r)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row[: len(cells[0])]):
            cell = t.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(re.sub(r"[*`]", "", val))
            run.font.size = Pt(9)
            run.bold = i == 0
    doc.add_paragraph()


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = MONO
            r.font.size = Pt(9.5)
        elif tok.startswith("*") and tok.endswith("*"):
            p.add_run(tok[1:-1]).italic = True
        else:
            p.add_run(tok)
    return p


def render(md, doc):
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            add_code(doc, block)
        elif line.strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            continue
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 2")
        elif line.startswith("# "):
            h = doc.add_paragraph(line[2:], style="Heading 1")
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("> "):
            p = add_para(doc, line[2:])
            p.paragraph_format.left_indent = Inches(0.35)
            p.runs and setattr(p.runs[0].font, "italic", True)
        elif re.match(r"^\s*[-*] ", line):
            add_para(doc, re.sub(r"^\s*[-*] ", "", line), style="List Bullet")
        elif re.match(r"^\s*\d+\. ", line):
            add_para(doc, re.sub(r"^\s*\d+\. ", "", line), style="List Number")
        elif line.startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
            if m:
                path = ROOT / m.group(2)
                if path.exists():
                    doc.add_picture(str(path), width=Inches(6.0))
                    cap = add_para(doc, m.group(1))
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs and setattr(cap.runs[0].font, "size", Pt(9))
                else:
                    add_para(doc, "[figure not yet generated: %s]" % m.group(2))
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip():
            add_para(doc, line)
        i += 1


def main():
    if not SRC.exists():
        sys.exit("no source at %s" % SRC)
    live = load_live()
    md, missing = substitute(SRC.read_text(), live)
    doc = Document()
    style_doc(doc)
    render(md, doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("wrote %s (%d live values, %d pending)" % (OUT, len(live), len(missing)))
    if missing:
        print("pending: " + ", ".join(sorted(missing)))


if __name__ == "__main__":
    main()
